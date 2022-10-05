from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class ResumeDoc:
    def __init__(self, file_root: str):
        self.file_root = file_root
        self.document = Document()

        font = self.document.styles['Heading 1'].font
        font.name = "Helvetica"
        font.size = Pt(14)
        font.color.rgb = RGBColor(31, 121, 199)
        font.bold = True
        self.document.styles['Heading 1'].paragraph_format.line_spacing = 1.5

        font = self.document.styles['Normal'].font
        font.name = "Helvetica"
        font.size = Pt(10)

        font = self.document.styles['Title'].font
        font.name = "Helvetica"
        font.size = Pt(42)
        font.bold = True
        font.color.rgb = RGBColor(5, 5, 5)

        font = self.document.styles['Body Text'].font
        font.name = "Helvetica"
        font.size = Pt(10)
        font.color.rgb = RGBColor(100, 100, 100)

        font = self.document.styles['Body Text 2'].font
        font.name = "Helvetica"
        font.size = Pt(14)
        self.document.styles['Body Text 2'].paragraph_format.line_spacing = 1

        style = self.document.styles['Body Text 3']
        font = style.font
        font.name = "Helvetica"
        font.size = Pt(10)
        font.bold = True
        style.paragraph_format.line_spacing = 1
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_header(self, resume):
        self.document.add_heading(resume["name"],
                                  0).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # self.document.add_paragraph(resume["blurb"])

    def add_experience(self, resume):
        self.section_header("EXPERIENCE")
        for experience in resume["experience"]:
            self.three_part(
                experience["company"], experience["location"], experience["title"]
            )
            self.document.add_paragraph(experience["dates"]).style = self.document.styles[
                'Body Text']
            if "description" in experience:
                self.document.add_paragraph(experience["description"])

    def add_education(self, resume):
        self.section_header("EDUCATION")
        for education in resume["education"]:
            self.three_part(education["school"], education["location"], education["degree"])
            self.document.add_paragraph(education["dates"]).style = self.document.styles[
                'Body Text']

    def add_contact(self, resume):
        self.document.add_paragraph(
            f"{resume['email']} * {resume['phone']} * {resume['address']}").style = \
            self.document.styles['Body Text 3']

    def add_skills(self, resume):
        self.document.add_heading("SKILLS")
        self.document.add_paragraph(" • ".join(resume["skills"]))

    def write_file(self, resume):
        self.add_header(resume)
        self.add_contact(resume)
        self.add_skills(resume)
        self.add_experience(resume)
        self.add_education(resume)
        self.document.save(f"{self.file_root}.docx")

    def section_header(self, name):
        self.document.add_heading(name, 1)

    def three_part(self, param1, param2, param3):
        p = self.document.add_paragraph("")
        p.style = self.document.styles['Body Text 2']
        run = p.add_run(param1)
        run.bold = True
        p.add_run(f", {param2}, ")
        p.add_run(param3).italic = True
