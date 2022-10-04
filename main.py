import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from fpdf import FPDF


class ResumeDoc:
    def __init__(self, file_root: str):
        self.file_root = file_root
        self.document = Document()

        font = self.document.styles['Normal'].font
        font.name = "Helvetica"
        font.size = Pt(10)

        font = self.document.styles['Title'].font
        font.name = "Helvetica"
        font.size = Pt(36)

        font = self.document.styles['Body Text'].font
        font.name = "Helvetica"
        font.size = Pt(9)
        font.color.rgb = RGBColor(100, 100, 100)

        font = self.document.styles['Body Text 2'].font
        font.name = "Helvetica"
        font.size = Pt(14)
        self.document.styles['Body Text 2'].paragraph_format.line_spacing = 1

        style = self.document.styles['Body Text 3']
        font = style.font
        font.name = "Helvetica"
        font.size = Pt(10)
        style.paragraph_format.line_spacing = 1
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_header(self, resume):
        self.document.add_heading(resume["name"],
                                  0).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # self.document.add_paragraph(resume["blurb"])

    def add_experience(self, resume):
        self.section_header("Experience")
        for experience in resume["experience"]:
            self.three_part(
                experience["company"], experience["location"], experience["title"]
            )
            self.document.add_paragraph(experience["dates"]).style = self.document.styles[
                'Body Text']
            if "description" in experience:
                self.document.add_paragraph(experience["description"])

    def add_education(self, resume):
        self.section_header("Education")
        for education in resume["education"]:
            self.three_part(education["school"], education["location"], education["degree"])
            self.document.add_paragraph(education["dates"])

    def add_contact(self, resume):

        self.document.add_paragraph(
            f"{resume['email']} * {resume['phone']} * {resume['address']}").style = \
        self.document.styles['Body Text 3']

    def add_skills(self, resume):
        self.document.add_heading("Skills")
        self.document.add_paragraph(" • ".join(resume["skills"]))


    def write_file(self, resume):
        self.add_header(resume)
        self.add_contact(resume)
        self.add_skills(resume)
        self.add_experience(resume)
        self.add_education(resume)
        self.document.save(f"{self.file_root}.docx")

    def section_header(self, name):
        self.document.add_heading(name, 1)

    def three_part(self, param1, param2, param3):
        p = self.document.add_paragraph("")
        p.style = self.document.styles['Body Text 2']
        run = p.add_run(param1)
        run.bold = True
        p.add_run(f", {param2}, ")
        p.add_run(param3).italic = True


class ResumePDF(FPDF):
    def __init__(self, file_root: str):
        super().__init__(format="letter")
        self.file_root = file_root

    def basic_bold(self):
        self.set_text_color(0, 0, 0)
        self.set_font("Helvetica", "B", 9)

    def title_emphasis(self, style=""):
        self.set_font("Helvetica", style, 12)

    def header_text(self, text):
        self.set_font("Helvetica", "B", 36)
        self.cell(125, 31, text, align="L")

    def lesser_text(self):
        self.set_text_color(100, 100, 100)
        self.set_font("Helvetica", "", 9)

    def section_header(self, value):
        self.ln()
        self.basic_bold()
        self.set_text_color(31, 121, 199)
        self.cell(self.get_string_width(value), self.font_size_pt, value.upper())
        self.ln()

    def three_part(self, first, second, third):
        self.set_text_color(0, 0, 0)
        self.title_emphasis("B")
        cell_width = self.get_string_width(first)
        self.cell(cell_width, 12, first)
        self.title_emphasis("")
        middle = f", {second} - "
        cell_width = self.get_string_width(middle)
        self.cell(cell_width, 12, middle)
        self.title_emphasis("I")
        cell_width = self.get_string_width(third)
        self.cell(cell_width, 12, third)

    def text_cell(self, text, height=None, lf=True):
        self.cell(
            self.get_string_width(text), height if height else self.font_size_pt, text
        )
        if lf:
            self.ln()

    def add_header(self, resume):
        self.header_text(resume["name"])
        self.set_font("Helvetica", "", 9)
        self.set_y(31)
        self.cell(125, 9, resume["blurb"], align="L")
        self.ln()

    def add_experience(self, resume):
        self.section_header("Experience")
        for experience in resume["experience"][0:3]:
            self.three_part(
                experience["company"], experience["location"], experience["title"]
            )
            self.set_y(self.get_y() + 7)
            self.lesser_text()
            self.text_cell(experience["dates"])
            self.multi_cell(125, 5, experience["description"])

    def add_education(self, resume):
        self.section_header("Education")
        for education in resume["education"]:
            self.three_part(education["school"], education["location"], education["degree"])
            self.set_y(self.get_y() + 7)
            self.lesser_text()
            self.text_cell(education["dates"])

    def add_additional_experience(self, resume):
        additional_experience = resume["experience"][3:]
        if not additional_experience:
            return
        self.section_header("Additional Work Experience")
        for experience in additional_experience:
            self.basic_bold()
            cell_width = self.get_string_width(experience["company"])
            self.cell(cell_width, 5, experience["company"])
            self.lesser_text()
            right = f" - {experience['location']} - {experience['title']} - {experience['dates']}"
            cell_width = self.get_string_width(right)
            self.cell(cell_width, 5, right)
            self.ln()

    def add_contact(self, resume):
        self.set_left_margin(150)
        self.set_y(15)
        self.basic_bold()
        self.text_cell(resume["phone"], height=5)
        self.text_cell(resume["email"], height=5)
        self.text_cell(resume["website"], height=5)
        self.text_cell(resume["address"], height=5)

    def add_skills(self, resume):
        self.set_y(45)
        self.section_header("Skills")
        self.lesser_text()
        for skill in resume["skills"]:
            self.text_cell(skill)

    def add_projects(self, resume):
        if not resume.get("projects"):
            return
        self.section_header("Projects")
        for project in resume["projects"]:
            self.basic_bold()
            self.text_cell(project["name"])
            self.lesser_text()
            self.multi_cell(50, 5, project["description"], align="L")
            self.ln()

    def write_file(self, resume):
        self.add_page()
        self.add_header(resume)
        self.add_experience(resume)
        self.add_additional_experience(resume)
        self.add_education(resume)
        self.add_contact(resume)
        self.add_skills(resume)
        self.add_projects(resume)
        self.output(f"{self.file_root}.pdf")


def write_resume(document):
    with open("resume.yml", "r") as resume_file:
        resume = yaml.safe_load(resume_file)
        document.write_file(resume)


if __name__ == "__main__":
    write_resume(ResumePDF("AJVResume"))
    write_resume(ResumeDoc("AJVResume"))
